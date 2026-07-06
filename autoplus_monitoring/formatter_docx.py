# -*- coding: utf-8 -*-
"""
formatter_docx.py
===============================================================================
[모듈 목적]
python-docx로 최종 모니터링 Word 문서를 생성한다.

[문서 구성 순서 - 정확히 지켜야 함]
  0. [경고] 부정 이슈 후보/VIG 민감 기사 (해당 시만)
  1. 오늘의 배포 자료 게재 현황 (배포 당일만)
  2. [오토플러스 뉴스] 자사 언급 기사
  3. [경쟁사 뉴스] 경쟁사별 섹션 구분
  4. [업계 뉴스] 우선순위 1→5 순 정렬

[기사 표기 형식]
"▷ [매체명] 기자명 | 제목 | URL" 다음 줄에 AI 판단 근거 1줄을 회색 텍스트로 표시.

[파일명]
output/모니터링_{YYYYMMDD}.docx

[사후 처리]
생성 완료 후 data/sent_articles.json 에 URL·제목·발송일을 자동 갱신한다.
===============================================================================
"""
import os
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from keywords import COMPETITOR_KEYWORDS, MEDIA_PRIORITY_ORDER
from release_calendar import get_today_releases, insert_release_section
from rule_filter import _now_kst_naive

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
SENT_ARTICLES_PATH = os.path.join(DATA_DIR, "sent_articles.json")

GRAY_COLOR = RGBColor(0x80, 0x80, 0x80)


def _add_hyperlink(paragraph, url, text, font_size_pt=10.5, color="0563C1", underline=True):
    """
    python-docx는 하이퍼링크를 기본 지원하지 않으므로, OOXML(w:hyperlink)을
    직접 구성해 문단에 클릭 가능한 하이퍼링크 run을 추가한다.
    (2026-07-06 추가: 기사 제목에 URL을 직접 걸기 위한 용도)
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size_pt * 2)))  # docx는 half-point 단위
    rPr.append(sz)

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_article_paragraph(doc, article: dict):
    """
    문서에 기사 한 건을 다음 형식으로 추가한다:
      ▷ [매체명] 기자명 | 제목(하이퍼링크로 URL 연결)
      (회색) AI 판단 근거 1줄

    [2026-07-06 변경] 기존에는 "제목 | URL"처럼 제목과 URL을 각각 텍스트로
    나열했으나, 담당자가 Word 문서에서 제목을 클릭해 바로 기사로 이동할 수
    있도록 제목 자체에 하이퍼링크를 거는 방식으로 변경했다. URL은 하이퍼링크에
    포함되므로 더 이상 별도 텍스트로 표시하지 않는다.
    """
    source = article.get("source", "미확인매체")
    journalist = article.get("journalist", "")
    title = article.get("title", "")
    url = article.get("url", "")
    reason = article.get("ai_reason", "")

    p = doc.add_paragraph()
    prefix_run = p.add_run(f"▷ [{source}] {journalist} | ")
    prefix_run.font.size = Pt(10.5)

    if url:
        _add_hyperlink(p, url, title, font_size_pt=10.5)
    else:
        # URL이 없는 경우(드묾) 안전하게 일반 텍스트로만 표시
        title_run = p.add_run(title)
        title_run.font.size = Pt(10.5)

    if reason:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"    - AI 판단 근거: {reason}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY_COLOR
        run2.italic = True


def _add_warning_section(doc, negative_flagged: list, vig_sensitive: list):
    """0. [경고] 부정 이슈 후보 / VIG 민감 기사 섹션 (해당 시만 표시)"""
    if not negative_flagged and not vig_sensitive:
        return

    heading = doc.add_heading("[경고] 부정 이슈 후보 / 민감 이슈", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    if negative_flagged:
        doc.add_heading("부정 키워드 포함 기사 (NEGATIVE_FLAG)", level=2)
        for article in negative_flagged:
            _add_article_paragraph(doc, article)

    if vig_sensitive:
        doc.add_heading("VIG파트너스 언급 기사 (실무 확인 필요)", level=2)
        for article in vig_sensitive:
            _add_article_paragraph(doc, article)
            note = article.get("sensitive_note", "")
            if note:
                p = doc.add_paragraph()
                run = p.add_run(f"    - {note}")
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.italic = True


def _add_own_news_section(doc, own_articles: list):
    """2. [오토플러스 뉴스] 자사 언급 기사"""
    doc.add_heading("[오토플러스 뉴스]", level=1)
    if not own_articles:
        doc.add_paragraph("해당 기간 내 자사 언급 기사가 없습니다.")
        return
    for article in own_articles:
        _add_article_paragraph(doc, article)


def _add_competitor_news_section(doc, competitor_articles: list):
    """3. [경쟁사 뉴스] 경쟁사별 섹션 구분"""
    doc.add_heading("[경쟁사 뉴스]", level=1)
    if not competitor_articles:
        doc.add_paragraph("해당 기간 내 경쟁사 관련 기사가 없습니다.")
        return

    # 경쟁사명별로 그룹핑 (search_keyword 기준, COMPETITOR_KEYWORDS 순서 유지)
    grouped = {kw: [] for kw in COMPETITOR_KEYWORDS}
    others = []
    for article in competitor_articles:
        kw = article.get("search_keyword", "")
        if kw in grouped:
            grouped[kw].append(article)
        else:
            others.append(article)

    for kw in COMPETITOR_KEYWORDS:
        articles = grouped[kw]
        if not articles:
            continue
        doc.add_heading(kw, level=2)
        for article in articles:
            _add_article_paragraph(doc, article)

    if others:
        doc.add_heading("기타 경쟁사 관련", level=2)
        for article in others:
            _add_article_paragraph(doc, article)


def _add_industry_news_section(doc, industry_articles: list):
    """4. [업계 뉴스] 우선순위 1→5 순 정렬"""
    doc.add_heading("[업계 뉴스]", level=1)
    if not industry_articles:
        doc.add_paragraph("해당 기간 내 업계 관련 기사가 없습니다.")
        return

    priority_labels = {
        "업계1": "1순위 - 중고차/렌터카/구독서비스",
        "업계2": "2순위 - 자동차금융(캐피탈/카드사)",
        "업계3": "3순위 - 브랜드 뉴스(완성차)",
        "업계4": "4순위 - 업계 기획/트렌드",
        "업계5": "5순위 - 기타(인프라/법규/자율주행)",
    }

    grouped = {k: [] for k in priority_labels}
    for article in industry_articles:
        cat = article.get("keyword_category", "")
        if cat in grouped:
            grouped[cat].append(article)

    for cat in ["업계1", "업계2", "업계3", "업계4", "업계5"]:
        articles = grouped[cat]
        if not articles:
            continue
        doc.add_heading(priority_labels[cat], level=2)
        for article in articles:
            _add_article_paragraph(doc, article)


def generate_monitoring_docx(ai_scored_results: dict, negative_flagged: list, vig_sensitive: list,
                              output_dir: str = OUTPUT_DIR):
    """
    최종 모니터링 Word 문서를 생성한다.

    Args:
        ai_scored_results: ai_scorer.run_ai_scoring() 반환값 중 "포함" 리스트를
            사용 (제외/검토필요는 문서에 포함하지 않되, 검토필요는 review_mode에서
            별도 취급)
        negative_flagged: rule_filter.apply_all_filters() 반환값의
            negative_flagged (통과분 내 서브셋)
        vig_sensitive: rule_filter.apply_all_filters() 반환값의 vig_sensitive

    Returns:
        생성된 문서의 절대 경로
    """
    os.makedirs(output_dir, exist_ok=True)

    included_articles = ai_scored_results.get("포함", [])

    own_articles = [a for a in included_articles if a.get("keyword_category") == "자사"]
    competitor_articles = [a for a in included_articles if a.get("keyword_category") == "경쟁사"]
    industry_articles = [a for a in included_articles if str(a.get("keyword_category", "")).startswith("업계")]

    doc = Document()

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(f"오토플러스 뉴스 모니터링 ({_now_kst_naive().strftime('%Y-%m-%d')})")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # 0. 경고 섹션 (부정이슈/VIG민감, 해당 시만)
    _add_warning_section(doc, negative_flagged, vig_sensitive)

    # 1. 오늘의 배포 자료 게재 현황 (배포 당일만)
    today_releases = get_today_releases()
    if today_releases:
        insert_release_section(doc, today_releases)

    # 2. 오토플러스 뉴스
    _add_own_news_section(doc, own_articles)

    # 3. 경쟁사 뉴스
    _add_competitor_news_section(doc, competitor_articles)

    # 4. 업계 뉴스
    _add_industry_news_section(doc, industry_articles)

    filename = f"모니터링_{_now_kst_naive().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"[formatter_docx] 문서 생성 완료 → {filepath}")

    # 발송 완료 처리: sent_articles.json 갱신
    _update_sent_articles(included_articles)

    return filepath


def _update_sent_articles(articles: list):
    """문서 생성 완료 후 data/sent_articles.json 에 URL·제목·발송일을 갱신한다."""
    if os.path.exists(SENT_ARTICLES_PATH):
        with open(SENT_ARTICLES_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = {"_comment": "기보고(발송 완료) 기사 URL/제목 이력.", "sent_articles": []}

    today_str = _now_kst_naive().strftime("%Y-%m-%d")
    existing_urls = {item.get("url") for item in data.get("sent_articles", [])}

    added = 0
    for article in articles:
        url = article.get("url", "")
        if url and url not in existing_urls:
            data["sent_articles"].append({
                "url": url,
                "title": article.get("title", ""),
                "sent_date": today_str,
            })
            existing_urls.add(url)
            added += 1

    with open(SENT_ARTICLES_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"[formatter_docx] sent_articles.json 갱신 완료 (신규 {added}건 추가)")


if __name__ == "__main__":
    dummy_results = {
        "포함": [
            {"title": "리본카, 신규 서비스 출시", "source": "이데일리", "journalist": "김기자",
             "url": "https://example.com/1", "keyword_category": "자사", "ai_reason": "자사 직접 언급"},
        ],
        "제외": [],
        "검토필요": [],
    }
    generate_monitoring_docx(dummy_results, [], [])
