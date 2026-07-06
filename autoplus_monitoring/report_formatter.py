# -*- coding: utf-8 -*-
"""
report_formatter.py
===============================================================================
[모듈 목적]
게재보고(당일 배포한 보도자료가 어느 매체에 게재되었는지 집계) 양식을 생성한다.
인수인계서 명시: "배포 당일 퇴근 전 커버리지 취합하여 실무 카톡방에 게재보고
진행."

경제지/무가지/오토지/온라인 구분 게재보고 표를 별도 함수로 생성한다
(MEDIA_PRIORITY_ORDER, media_normalizer의 media_group 필드 활용).
===============================================================================
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt

from keywords import MEDIA_PRIORITY_ORDER

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")

# 게재보고에서 강조하는 4개 구분 (인수인계서 "경제지/무가지/오토지/온라인 구분" 명시)
REPORT_GROUPS = ["경제지", "전문지/무가지", "오토지", "온라인"]


def build_coverage_table(articles: list, release_title: str = ""):
    """
    특정 보도자료(release_title)에 대한 게재 현황을 매체그룹별로 집계한다.
    articles는 media_normalizer.normalize_articles() 처리가 완료되어
    media_group 필드가 채워져 있어야 한다.

    반환값: {"경제지": [...], "전문지/무가지": [...], "오토지": [...], "온라인": [...], "기타": [...]}
    """
    grouped = {g: [] for g in REPORT_GROUPS}
    grouped["기타"] = []

    for article in articles:
        group = article.get("media_group", "온라인")
        if group in grouped:
            grouped[group].append(article)
        else:
            grouped["기타"].append(article)

    total = sum(len(v) for v in grouped.values())
    print(f"[report_formatter] 게재보고 집계 완료 - 총 {total}건 "
          f"(경제지 {len(grouped['경제지'])}, 전문지/무가지 {len(grouped['전문지/무가지'])}, "
          f"오토지 {len(grouped['오토지'])}, 온라인 {len(grouped['온라인'])}, 기타 {len(grouped['기타'])})")

    return grouped


def generate_coverage_report_docx(articles: list, release_title: str = "", output_dir: str = OUTPUT_DIR):
    """
    게재보고 Word 문서를 생성한다. 표 형식으로 매체그룹/매체명/제목/URL을 정리한다.
    """
    os.makedirs(output_dir, exist_ok=True)

    grouped = build_coverage_table(articles, release_title)
    total = sum(len(v) for v in grouped.values())

    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run(f"게재보고 - {release_title or '(제목 미지정)'}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"총 게재 건수: {total}건 (집계일: {datetime.now().strftime('%Y-%m-%d')})")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "구분"
    hdr_cells[1].text = "매체명"
    hdr_cells[2].text = "제목"
    hdr_cells[3].text = "URL"

    for group in REPORT_GROUPS + ["기타"]:
        for article in grouped[group]:
            row_cells = table.add_row().cells
            row_cells[0].text = group
            row_cells[1].text = article.get("source", "")
            row_cells[2].text = article.get("title", "")
            row_cells[3].text = article.get("url", "")

    filename = f"게재보고_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"[report_formatter] 게재보고 문서 생성 완료 → {filepath}")
    return filepath


if __name__ == "__main__":
    sample = [
        {"source": "이데일리", "title": "테스트 기사1", "url": "https://example.com/1", "media_group": "경제지"},
        {"source": "오토IN", "title": "테스트 기사2", "url": "https://example.com/2", "media_group": "오토지"},
    ]
    generate_coverage_report_docx(sample, "테스트 보도자료")
